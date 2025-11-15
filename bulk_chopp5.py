import os
import sys
import argparse
from docx import Document
from google.cloud import vision

# Supported image formats
SUPPORTED_FORMATS = ('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.webp')

def initialize_client(credentials_path):
    """Initialize Google Vision client with provided credentials."""
    if not os.path.exists(credentials_path):
        print(f"Error: Credentials file not found at {credentials_path}")
        sys.exit(1)
    
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = credentials_path
    try:
        return vision.ImageAnnotatorClient()
    except Exception as e:
        print(f"Error initializing Google Vision client: {e}")
        sys.exit(1)

def extract_text_from_image(image_path, client):
    """Extract text from an image using Google Vision API."""
    try:
        with open(image_path, "rb") as image_file:
            content = image_file.read()

        image = vision.Image(content=content)
        response = client.text_detection(image=image)
        
        if response.error.message:
            print(f"Warning: API error for {image_path}: {response.error.message}")
            return ""
        
        texts = response.text_annotations
        if texts:
            return texts[0].description
        return ""
    except Exception as e:
        print(f"Warning: Failed to process {image_path}: {e}")
        return ""

def process_folders_to_docx(main_folder_path, output_docx, client):
    """Process a folder or subfolders and save extracted raw text to a Word document."""
    if not os.path.exists(main_folder_path):
        print(f"Error: Input folder not found at {main_folder_path}")
        sys.exit(1)
    
    document = Document()
    document.add_heading("Extracted Text Data", level=1)

    # Check if the path contains image files directly
    contains_files = any(
        filename.lower().endswith(SUPPORTED_FORMATS)
        for filename in os.listdir(main_folder_path)
    )

    if contains_files:
        # Process each image in the folder individually
        print(f"Processing images in main folder...")
        for filename in sorted(os.listdir(main_folder_path)):
            if filename.lower().endswith(SUPPORTED_FORMATS):
                image_path = os.path.join(main_folder_path, filename)
                document.add_heading(f"Image: {filename}", level=2)

                # Extract text from the image
                text = extract_text_from_image(image_path, client)

                # Add raw text to the document
                document.add_paragraph(text if text else "[No text detected]")
                print(f"  [OK] Processed: {filename}")
    else:
        # Process subfolders in the main folder
        subdirs = sorted([d for d in os.listdir(main_folder_path) 
                         if os.path.isdir(os.path.join(main_folder_path, d))])
        
        if not subdirs:
            print("Warning: No images or subdirectories found in the specified folder.")
            sys.exit(1)
        
        for idx, subfolder in enumerate(subdirs, 1):
            subfolder_path = os.path.join(main_folder_path, subfolder)
            document.add_heading(f"Directory {idx}: {subfolder}", level=2)
            print(f"Processing Directory {idx}: {subfolder}")

            images_found = False
            for filename in sorted(os.listdir(subfolder_path)):
                if filename.lower().endswith(SUPPORTED_FORMATS):
                    images_found = True
                    image_path = os.path.join(subfolder_path, filename)
                    document.add_heading(f"File: {filename}", level=3)

                    # Extract text from the image
                    text = extract_text_from_image(image_path, client)

                    # Add raw text to the document
                    document.add_paragraph(text if text else "[No text detected]")
                    print(f"  [OK] Processed: {filename}")
            
            if not images_found:
                document.add_paragraph("[No images found in this directory]")
                print(f"  [WARN] No images found in {subfolder}")

    # Save the Word document
    try:
        document.save(output_docx)
        print(f"\n[OK] Results successfully saved to {output_docx}")
    except Exception as e:
        print(f"Error: Failed to save document: {e}")
        sys.exit(1)

def main():
    """Main CLI entry point."""
    parser = argparse.ArgumentParser(
        description="Extract text from images using Google Cloud Vision API and save to Word document.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s -c credentials.json -i ./images -o output.docx
  %(prog)s --credentials /path/to/creds.json --input ./my_folder --output results.docx
        """
    )
    
    parser.add_argument(
        "-c", "--credentials",
        required=True,
        help="Path to Google Cloud Vision API credentials JSON file"
    )
    
    parser.add_argument(
        "-i", "--input",
        required=True,
        help="Path to folder containing images or subfolders with images"
    )
    
    parser.add_argument(
        "-o", "--output",
        required=True,
        help="Path for output Word document (e.g., results.docx)"
    )
    
    args = parser.parse_args()
    
    # Initialize client
    print("Initializing Google Cloud Vision client...")
    client = initialize_client(args.credentials)
    
    # Process folders
    process_folders_to_docx(args.input, args.output, client)

if __name__ == "__main__":
    main()
